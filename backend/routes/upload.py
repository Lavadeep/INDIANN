from fastapi import APIRouter, UploadFile, Form, Depends, HTTPException, status
from sqlalchemy.orm import Session
from datetime import datetime
import io
import re
import chardet
import nltk

from langdetect import detect
from indicnlp.tokenize import sentence_tokenize

from .. import models, database
from .permissions import get_current_user, require_project_role_by_project

# Optional: PDF/DOCX extraction (install PyPDF2, python-docx if needed)
try:
    import PyPDF2
except ImportError:
    PyPDF2 = None
try:
    import docx
except ImportError:
    docx = None

router = APIRouter()

nltk.download("punkt", quiet=True)

LANGUAGE_ALIASES = {
    # English
    "en": "en",
    "eng": "en",
    "english": "en",
    # Hindi
    "hi": "hi",
    "hin": "hi",
    "hindi": "hi",
    # Bengali
    "bn": "bn",
    "ben": "bn",
    "bengali": "bn",
    # Odia / Oriya
    "or": "or",
    "ori": "or",
    "odia": "or",
    "oriya": "or",
    # Telugu
    "te": "te",
    "tel": "te",
    "telugu": "te",
    # Tamil
    "ta": "ta",
    "tam": "ta",
    "tamil": "ta",
    # Kannada
    "kn": "kn",
    "kan": "kn",
    "kannada": "kn",
    # Malayalam
    "ml": "ml",
    "mal": "ml",
    "malayalam": "ml",
    # Gujarati
    "gu": "gu",
    "guj": "gu",
    "gujarati": "gu",
    # Marathi
    "mr": "mr",
    "mar": "mr",
    "marathi": "mr",
}

INDIC_SENTENCE_TOKENIZE_LANGS = {
    "hi",
    "bn",
    "or",
    "te",
    "ta",
    "kn",
    "ml",
    "gu",
    "mr",
}


def normalize_language_code(lang: str) -> str:
    key = (lang or "").lower().strip()
    if not key:
        return ""
    return LANGUAGE_ALIASES.get(key, key[:2])


def protect_decimal_points(text: str) -> str:
    # Prevent sentence splitters from breaking numeric decimals like 0.3 / 3.3 / 2.5
    # by temporarily replacing the dot between digits.
    return re.sub(r"(?<=\d)\.(?=\d)", "<DECIMAL_DOT>", text)


def restore_decimal_points(text: str) -> str:
    return text.replace("<DECIMAL_DOT>", ".")


# ------------------------------------------------------------
# Normalize plain text for non-ConLLU uploads (PDF, DOCX, TXT, etc.):
# - Whitespace normalization (single spaces, strip)
# - Sentence segmentation (nltk for English, Indic NLP for Indian languages)
# - Store as one sentence per line so frontend can use content.split("\n")
# All offsets are computed against this finalized stored content.
# ------------------------------------------------------------
def normalize_text(text: str, lang: str) -> str:
    text = re.sub(r"\s+", " ", text).strip()
    if not text:
        return ""
    text = protect_decimal_points(text)
    lang_code = normalize_language_code(lang)
    try:
        if lang_code in INDIC_SENTENCE_TOKENIZE_LANGS:
            sents = sentence_tokenize.sentence_split(text, lang_code)
        else:
            sents = nltk.sent_tokenize(text)
    except Exception:
        sents = [text]
    normalized = "\n".join(s.strip() for s in sents if s.strip())
    return restore_decimal_points(normalized)


def extract_text_from_binary(raw_bytes: bytes, file_type: str) -> str:
    """
    Extract plain text from PDF or DOCX binary content.
    For other types, returns None (caller should decode as text).
    """
    ft = file_type.lower()
    if ft == "pdf":
        if PyPDF2 is None:
            raise HTTPException(
                status_code=status.HTTP_501_NOT_IMPLEMENTED,
                detail="PDF extraction requires PyPDF2. Install with: pip install PyPDF2",
            )
        try:
            reader = PyPDF2.PdfReader(io.BytesIO(raw_bytes))
            parts = []
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    parts.append(t)
            return "\n".join(parts) if parts else ""
        except Exception as e:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Failed to extract text from PDF: {e!s}",
            )
    if ft in ("docx", "word"):
        if docx is None:
            raise HTTPException(
                status_code=status.HTTP_501_NOT_IMPLEMENTED,
                detail="DOCX extraction requires python-docx. Install with: pip install python-docx",
            )
        try:
            doc = docx.Document(io.BytesIO(raw_bytes))
            return "\n".join(p.text for p in doc.paragraphs)
        except Exception as e:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Failed to extract text from DOCX: {e!s}",
            )
    return None


# ------------------------------------------------------------
# Upload Endpoint
# ------------------------------------------------------------
@router.post("/upload")
async def upload_file(
    project_id: int = Form(...),
    uploaded_by: int = Form(...),
    file_type: str = Form(...),
    file: UploadFile = Form(...),
    db: Session = Depends(database.get_db),
    current_user: models.User = Depends(get_current_user),
):
    """
    Upload document.

    - Only project owner/admin can upload
    - Curators are forbidden
    - For .conllu:
        * store RAW CoNLL-U text
        * auto-create POS + dependency annotations
    """

    # ---------------- Permission checks ----------------
    if uploaded_by != current_user.user_id:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="uploaded_by must match authenticated user",
        )

    role = require_project_role_by_project(
        project_id=project_id, db=db, current_user=current_user
    )

    if role == "curator":
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Curators are not allowed to upload documents",
        )

    if role not in ("owner", "admin"):
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Only project owners or admins can upload documents",
        )

    # ---------------- Read file ----------------
    raw_bytes = await file.read()
    ft = file_type.lower()

    # PDF/DOCX: extract text from binary; others: decode as text
    extracted = extract_text_from_binary(raw_bytes, ft)
    if extracted is not None:
        text = extracted
    else:
        detected = chardet.detect(raw_bytes)
        encoding = detected.get("encoding") or "utf-8"
        text = raw_bytes.decode(encoding, errors="ignore")

    # ---------------- Language detection ----------------
    try:
        detected_lang = detect(text)
    except Exception:
        detected_lang = "unknown"

    project = (
        db.query(models.Project)
        .filter(models.Project.project_id == project_id)
        .first()
    )
    if not project:
        raise HTTPException(400, "Invalid project_id")

    project_lang_raw = project.language or ""
    detected_lang_raw = detected_lang or ""
    project_lang = normalize_language_code(project_lang_raw)
    detected_lang = normalize_language_code(detected_lang_raw)
    # print(
    #     "Upload language detection:",
    #     {
    #         "project_id": project_id,
    #         "file": file.filename,
    #         "project_lang_raw": project_lang_raw,
    #         "project_lang": project_lang,
    #         "detected_lang_raw": detected_lang_raw,
    #         "detected_lang": detected_lang,
    #     },
    # )

    language_mismatch = False
    if project_lang and detected_lang != "unknown":
        if project_lang != detected_lang:
            language_mismatch = True

    # ---------------- Normalize content for non-ConLLU (PDF, DOCX, TXT, etc.) ----------------
    # CoNLL-U is already sentence-segmented (one sentence per block); store as-is.
    # Other formats: apply sentence segmentation + whitespace normalization so stored
    # content is "{sentence}\n{sentence}\n..."; frontend uses split("\n"). Offsets
    # are computed against this finalized content.
    ft = file_type.lower()
    if ft != "conllu":
        lang = project_lang or detected_lang or "en"
        text = normalize_text(text, lang)
        uploaded_sents = [s for s in text.split("\n") if s.strip()]
        print(
            "Sentences being uploaded:",
            {
                "document": file.filename,
                "count": len(uploaded_sents),
            },
        )
        for i, sent in enumerate(uploaded_sents, start=1):
            print(f"{i:03d}: {sent}")

    # ---------------- Store Document ----------------
    new_doc = models.Document(
        project_id=project_id,
        filename=file.filename,
        file_type=ft,
        content=text,  # Canonical: sentence-per-line for non-ConLLU; raw CoNLL-U for ConLLU
        uploaded_by=current_user.user_id,
        uploaded_at=datetime.utcnow(),
    )

    db.add(new_doc)
    db.commit()
    db.refresh(new_doc)

    # ============================================================
    # CoNLL-U: parse tokens and create POS + dependency annotations
    # ============================================================
    if ft == "conllu":
        lines = text.splitlines()

        for line in lines:
            line = line.strip()
            if not line or line.startswith("#"):
                continue

            parts = line.split("\t")
            if len(parts) < 10:
                continue

            (
                token_id,
                form,
                lemma,
                upos,
                xpos,
                feats,
                head,
                deprel,
                deps,
                misc,
            ) = parts

            # Skip ranges (1-2) and empty nodes (3.1)
            if "-" in token_id or "." in token_id:
                continue

            # -------- Extract offsets from MISC --------
            start = end = None
            for field in misc.split("|"):
                if field.startswith("start="):
                    try:
                        start = int(field.replace("start=", ""))
                    except ValueError:
                        pass
                elif field.startswith("end="):
                    try:
                        end = int(field.replace("end=", ""))
                    except ValueError:
                        pass

            # -------- POS (EntityAnnotation) --------
            # In CoNLL-U, "_" in UPOS means "no label"; do not create an annotation for it.
            if start is not None and end is not None and upos and upos != "_":
                pos_anno = models.EntityAnnotation(
                    document_id=new_doc.document_id,
                    user_id=current_user.user_id,
                    start_offset=start,
                    end_offset=end,
                    entity_label=upos,
                    entity_text=form,
                )
                db.add(pos_anno)

            # -------- DependencyAnnotation --------
            try:
                dep_anno = models.DependencyAnnotation(
                    document_id=new_doc.document_id,
                    token_index=int(token_id),
                    head_index=int(head),
                    deprel=deprel,
                    user_id=current_user.user_id,
                )
                db.add(dep_anno)
            except Exception:
                # Defensive: malformed head/token id
                pass

        db.commit()

    # ---------------- Done ----------------
    return {
        "status": "success",
        "document_id": new_doc.document_id,
        "language_mismatch": language_mismatch,
    }
